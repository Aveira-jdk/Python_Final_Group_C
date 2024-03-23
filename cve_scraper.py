"""
Dear Programmer:
When We wrote this code, only god, and
We knew how it worked.
Now, only god know it!

Therefore, if you are trying to optimize
this routine, and it fails (most surely),
please increase this counter as a
warning for the next person:

total hours wasted here = 75

"""


import os
import re
import json
import webbrowser
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF  # Import FPDF for PDF generation
import time  # Import time module for delay

# Set up Selenium WebDriver
driver = webdriver.Firefox()  # change this parameter if you have different browser

# Declaring url patterns for websites for further use
nvd_url_pattern = "https://nvd.nist.gov/vuln/detail/{cve_id}"
exploit_db_url_pattern = "https://www.exploit-db.com/search?cve={cve_id}"
cve_url_pattern = "https://www.cvedetails.com/cve/{cve_id}/"
vulmon_url_pattern = "https://vulmon.com/vulnerabilitydetails?qid={cve_id}&scoretype=cvssv3"


# nvd scraper by Mahammad
def get_nvd_info(cve_id):
    # Construct the URL using the CVE ID
    url = nvd_url_pattern.format(cve_id=cve_id)

    # Access the URL in a WebDriver
    driver.get(url)

    try:
        # Find the element containing vulnerability summary and wait for its presence
        summary_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//p[@data-testid='vuln-description']"))
        )
        # Extract the vulnerability summary
        summary = summary_element.text
    except:
        summary = "Summary is not available"

    try:
        # Find the element containing CVSS score
        cvss_score_element = driver.find_element(By.XPATH, "//a[@data-testid='vuln-cvss3-panel-score']")
        # Extract the CVSS score
        cvss_score = cvss_score_element.text
    except:
        cvss_score = "Base score is not available"

    try:
        # Find the element containing CVSS vector
        vector_element = driver.find_element(By.XPATH, "//span[@data-testid='vuln-cvss3-nist-vector']")
        # Extract the CVSS vector
        vector = vector_element.text
    except:
        vector = "Vector is not available"

    try:
        # Find the element containing reference link
        reference_element = driver.find_element(By.XPATH, "//td[@data-testid='vuln-hyperlinks-link-0']/a")
        # Extract the reference link
        reference = reference_element.get_attribute("href")
    except:
        reference = "Reference is not available"

    # Return the gathered information
    return summary, cvss_score, vector, reference


# affected systems scraper function by Mahammad
def affected_systems(cve_id):
    url = cve_url_pattern.format(cve_id=cve_id)
    driver.get(url)

    try:
        # Wait for the table to load
        table = driver.find_element(By.XPATH, "/html/body/div[1]/div/div[2]/div/main/div[7]/ul/li")
        # Get the text content of the table
        table_content = table.text
        return table_content
    except Exception as e:
        return f"An error occurred: {e}"


# exploit lists from Vulmon by Nuray
# def vulmon_exploit_list_desc(cve_id):
#     # Set Firefox options
#     firefox_options = webdriver.FirefoxOptions()
#     firefox_options.headless = True
#
#     # Set up Firefox WebDriver with headless mode
#     driver = webdriver.Firefox(options=firefox_options)
#
#     exploit_desc = ""
#
#     try:
#         # Construct the URL with the provided CVE ID
#         url = vulmon_url_pattern.format(cve_id=cve_id)
#
#         # Open the URL
#         driver.get(url)
#
#         # Wait for the content to be loaded
#         content_element = driver.find_element(By.XPATH, '/html/body/div[3]/div/div[1]/div[5]')
#
#         # Get the content
#         exploit_desc = content_element.text
#
#     except Exception as e:
#         print("Error:", e)
#     finally:
#         # Close the browser
#         driver.quit()
#
#     return exploit_desc


# exploit-db scraper by Ravan
def check_exploit_db(cve_id):
    # Set up Selenium Firefox driver
    driver = webdriver.Firefox()

    # Construct the URL based on the CVE ID
    url = f"https://www.exploit-db.com/search?cve={cve_id}"

    # Open the URL in the browser
    driver.get(url)

    # Wait for the page to load (let's improve this comment to explain why we're waiting)
    # Waiting for a fixed time might not be the most robust approach, consider alternatives like WebDriverWait
    time.sleep(5)

    # Find the table containing the links
    table = driver.find_element(By.XPATH, '//*[@id="exploits-table"]')

    # Find all links in the table
    links = table.find_elements(By.TAG_NAME, 'a')

    # Extract and return the href attribute of each link
    return [link.get_attribute('href') for link in links]


# docx export by Aytan
def create_docx_report(cve_id, nvd_info, affected_systems_info, exploit_links):  # add exploit_desc if you want extra
    summary, cvss_score, vector, reference = nvd_info

    doc = Document()

    # Title
    title = doc.add_heading('CVE Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CVE Information
    cve_paragraph = doc.add_paragraph()
    cve_paragraph.add_run("CVE ID: ").bold = True
    cve_paragraph.add_run(f"{cve_id}\n\n")

    cve_paragraph.add_run("Summary: ").bold = True
    cve_paragraph.add_run(f"{summary}\n\n")

    cve_paragraph.add_run("Base Score: ").bold = True
    cve_paragraph.add_run(f"{cvss_score}\n\n")

    cve_paragraph.add_run("Vector: ").bold = True
    cve_paragraph.add_run(f"{vector}\n\n")

    cve_paragraph.add_run("Reference: ").bold = True
    cve_paragraph.add_run(f"{reference}\n\n")

    # Affected systems information
    affected_sys_paragraph = doc.add_paragraph()
    affected_sys_paragraph.add_run("Affected Systems Information: ").bold = True
    affected_sys_paragraph.add_run(f"{affected_systems_info}\n\n")

    # Exploit Descriptions
    # exploit_desc_paragraph = doc.add_paragraph()
    # exploit_desc_paragraph.add_run("Exploit Descriptions:").bold = True
    # exploit_desc_paragraph.add_run(f"\n{exploit_desc}\n\n")

    # Exploit Information
    exploit_paragraph = doc.add_paragraph()

    download_exploits = [link for link in exploit_links if '/download/' in link]
    look_exploits = [link for link in exploit_links if '/download/' not in link]

    if download_exploits:
        exploit_paragraph.add_run("Exploit Links to Download:").bold = True
        exploit_paragraph.add_run("\n\n")
        for link in download_exploits:
            exploit_paragraph.add_run(f"{link}\n")

        exploit_paragraph.add_run("\n" * 2)

    if look_exploits:
        exploit_paragraph.add_run("Exploit Links to Look:").bold = True
        exploit_paragraph.add_run("\n\n")
        for link in look_exploits:
            exploit_paragraph.add_run(f"{link}\n")

    # Save the document
    folder_name = f"reports/{cve_id}"
    os.makedirs(folder_name, exist_ok=True)
    docx_filename = os.path.join(folder_name, f"report_{cve_id}.docx")
    doc.save(docx_filename)

    return docx_filename


# json export by Nuray
def export_to_json(cve_id, nvd_info, exploit_links, affected_systems_info):
    # Unpack NVD information
    summary, cvss_score, vector, reference = nvd_info

    # Separate exploit links into download and non-download links
    download_exploits = [link for link in exploit_links if '/download/' in link]
    look_exploits = [link for link in exploit_links if '/download/' not in link]

    # Prepare data for JSON export
    data = {
        "CVE_ID": cve_id,
        "NVD_Info": {
            "Summary": summary,
            "Base_Score": cvss_score,
            "Vector": vector,
            "Reference": reference
        },
        "Affected_Systems_Info": affected_systems_info,
        "Exploit_Links": {
            "Download_Exploits": download_exploits,
            "Look_Exploits": look_exploits
        }
    }

    # Create a folder to store reports if it doesn't exist
    folder_name = f"reports/{cve_id}"
    os.makedirs(folder_name, exist_ok=True)

    # Define the JSON file name
    json_filename = os.path.join(folder_name, f"{cve_id}.json")

    # Write data to JSON file
    with open(json_filename, "w") as json_file:
        json.dump(data, json_file, indent=4)

    # Return the path of the generated JSON file
    return json_filename


# export to PDF by Aytan
def export_to_pdf(cve_id, nvd_info, exploit_links, affected_systems_info):
    summary, cvss_score, vector, reference = nvd_info

    # Create PDF object
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Add a page
    pdf.add_page()

    # Set font for the entire document
    pdf.set_font("Arial", size=12)

    # Add title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, f"CVE Report for {cve_id}", ln=True, align='C')
    pdf.ln(10)

    # Add CVE Information section
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, txt="CVE Information", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="CVE ID: " + cve_id, ln=True)
    pdf.multi_cell(0, 5, txt="Summary: " + summary)
    pdf.cell(0, 10, txt="Base Score: " + cvss_score, ln=True)
    pdf.cell(0, 10, txt="Vector: " + vector, ln=True)

    # Modified Reference Section
    pdf.set_font("Arial", size=12)  # Smaller font for reference
    pdf.cell(0, 6, txt="Reference:", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 6, txt=reference, ln=True)

    # Add Affected Systems Information section
    pdf.ln(15)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, txt="Affected Systems Information", ln=True)
    pdf.ln(5)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 5, txt=affected_systems_info)

    # Add Exploit Information section
    pdf.ln(15)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, txt="Exploit Information", ln=True)
    pdf.ln(5)

    download_exploits = [link for link in exploit_links if '/download/' in link]
    look_exploits = [link for link in exploit_links if '/download/' not in link]

    if download_exploits:
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, txt="Exploit Links to Download:", ln=True)
        for link in download_exploits:
            pdf.multi_cell(0, 5, txt=link)

    if look_exploits:
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, txt="Exploit Links to Look:", ln=True)
        for link in look_exploits:
            pdf.multi_cell(0, 5, txt=link)

    # Save the PDF file
    folder_name = f"reports/{cve_id}"
    os.makedirs(folder_name, exist_ok=True)
    pdf_filename = os.path.join(folder_name, f"report_{cve_id}.pdf")
    pdf.output(pdf_filename)
    return pdf_filename


# export to md function by Ravan
def export_to_markdown(cve_id, nvd_info, exploit_links, affected_systems_info):
    summary, cvss_score, vector, reference = nvd_info

    # Construct markdown content
    markdown_content = f"# CVE Report for {cve_id}\n\n"
    markdown_content += "## CVE Information\n"
    markdown_content += f"- **CVE ID:** {cve_id}\n"
    markdown_content += f"- **Summary:** {summary}\n"
    markdown_content += f"- **Base Score:** {cvss_score}\n"
    markdown_content += f"- **Vector:** {vector}\n"
    markdown_content += f"- **Reference:** [{reference}]({reference})\n\n"

    markdown_content += "## Affected Systems Information\n"
    markdown_content += f"{affected_systems_info}\n\n"

    markdown_content += "## Exploit Information\n"
    download_exploits = [link for link in exploit_links if '/download/' in link]
    look_exploits = [link for link in exploit_links if '/download/' not in link]

    if download_exploits:
        markdown_content += "### Exploit Links to Download:\n"
        for link in download_exploits:
            markdown_content += f"- [{link}]({link})\n"
        markdown_content += "\n"

    if look_exploits:
        markdown_content += "### Exploit Links to Look:\n"
        for link in look_exploits:
            markdown_content += f"- [{link}]({link})\n"
        markdown_content += "\n"

    # Save the Markdown content to a file
    folder_name = f"reports/{cve_id}"
    os.makedirs(folder_name, exist_ok=True)
    markdown_filename = os.path.join(folder_name, f"report_{cve_id}.md")
    with open(markdown_filename, "w") as markdown_file:
        markdown_file.write(markdown_content)

    return markdown_filename


# export to webpage by Mahammad
def export_to_html(cve_id, nvd_info, exploit_links, affected_systems_info):
    # Unpack NVD information
    summary, cvss_score, vector, reference = nvd_info

    # Separate exploit links into download and non-download links
    download_exploits = [link for link in exploit_links if '/download/' in link]
    look_exploits = [link for link in exploit_links if '/download/' not in link]

    # HTML content template
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>CVE Report - {cve_id}</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-image: url('https://img2.wallspic.com/crops/9/0/1/9/6/169109/169109-anime_landscape-anime-landscape_painting-art-landscape-3840x2160.jpg');
                background-size: cover;
                background-position: center;
                color: #ffffff;
            }}
            h1 {{
                text-align: center;
                color: #ffffff;
                background-color: #5b00ff;
                padding: 10px;
            }}
            h2 {{
                margin-bottom: 5px;
                color: #ffffff;
            }}
            p {{
                margin-bottom: 5px;
            }}
            a {{
                color: #ffffff;
                text-decoration: none;
            }}
            a:hover {{
                text-decoration: underline;
            }}
            ul {{
                list-style-type: none;
                padding: 0;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        <h1>CVE Report for {cve_id}</h1>

        <h2>CVE Information</h2>
        <p><strong>CVE ID: </strong> {cve_id}</p>
        <p><strong>Summary: </strong> {summary} </p>
        <p><strong>Base Score: </strong> {cvss_score}</p>
        <p><strong>Vector: </strong> {vector}</p>
        <p><strong>Reference: </strong> <a href="{reference}">{reference}</a></p>

        <h2>Affected Systems Information</h2>
        <p>{affected_systems_info}</p>

        <h2>Exploit Information</h2>
        """

    # Add download exploit links if available
    if download_exploits:
        html_content += "<h3>Exploit Links to Download</h3><ul>"
        for link in download_exploits:
            html_content += f"<li><a href='{link}'>{link}</a></li>"
        html_content += "</ul>"

    # Add non-download exploit links if available
    if look_exploits:
        html_content += "<h3>Exploit Links to Look</h3><ul>"
        for link in look_exploits:
            html_content += f"<li><a href='{link}'>{link}</a></li>"
        html_content += "</ul>"

    # Close HTML content
    html_content += """
    </body>
    </html>
    """

    # Create folder to store reports if it doesn't exist
    folder_name = f"reports/{cve_id}"
    os.makedirs(folder_name, exist_ok=True)

    # Define HTML file name
    html_filename = os.path.join(folder_name, f"report_{cve_id}.html")

    # Write HTML content to file
    with open(html_filename, 'w') as f:
        f.write(html_content)

    # Return the path of the generated HTML file
    return html_filename


# tab opener by Mahammad
def open_references(cve_id):
    # Construct URLs based on the CVE ID
    nvd_url = nvd_url_pattern.format(cve_id=cve_id)
    exploit_db_url = exploit_db_url_pattern.format(cve_id=cve_id)
    cve_url = cve_url_pattern.format(cve_id=cve_id)
    vulmon_url = vulmon_url_pattern.format(cve_id=cve_id)

    # Open tabs for NVD, Exploit-DB, CVE Details, and Vulmon
    webbrowser.open_new_tab(nvd_url)
    webbrowser.open_new_tab(exploit_db_url)
    webbrowser.open_new_tab(cve_url)
    webbrowser.open_new_tab(vulmon_url)


# cve checker by Nuray
def cve_exists(cve_id):
    """
    Check if the given CVE ID exists by attempting to scrape information from relevant sources.
    If no valid information is obtained, it is assumed that the CVE does not exist.

    Args:
        cve_id (str): The CVE ID to check.

    Returns:
        bool: True if the CVE exists, False otherwise.
    """
    # Attempt to retrieve NVD information
    nvd_info = get_nvd_info(cve_id)
    if nvd_info[0] == "Summary is not available":
        return False

    # Attempt to retrieve affected systems information
    affected_systems_info = affected_systems(cve_id)
    if not affected_systems_info:
        return False

    # Attempt to retrieve exploit links
    exploit_links = check_exploit_db(cve_id)
    if not exploit_links:
        return False

    # If information is obtained from all sources, assume CVE exists
    return True


# main function by Mahammad
def main():
    while True:
        print("[1]: Continue")
        print("[2]: Exit")

        choice = input("Enter your choice: ")

        if choice == "1":
            cve_id = input("Enter CVE ID: ").strip()
            cve_regex = r'^CVE-\d{4}-\d{4,}$'
            if re.match(cve_regex, cve_id):
                # Check if the CVE exists
                if cve_exists(cve_id):
                    nvd_info = get_nvd_info(cve_id)
                    affected_systems_info = affected_systems(cve_id)
                    exploit_links = check_exploit_db(cve_id)

                    docx_filename = create_docx_report(cve_id, nvd_info, affected_systems_info, exploit_links)
                    json_filename = export_to_json(cve_id, nvd_info, exploit_links, affected_systems_info)
                    pdf_filename = export_to_pdf(cve_id, nvd_info, exploit_links, affected_systems_info)
                    markdown_filename = export_to_markdown(cve_id, nvd_info, exploit_links, affected_systems_info)
                    html_filename = export_to_html(cve_id, nvd_info, exploit_links, affected_systems_info)

                    print(f"DOCX report saved as: {docx_filename}")
                    print(f"JSON data exported as: {json_filename}")
                    print(f"PDF report saved as: {pdf_filename}")
                    print(f"Markdown report saved as: {markdown_filename}")
                    print(f"HTML report saved as: {html_filename}")

                    open_references(cve_id)
                else:
                    print("CVE not found. Please enter a valid CVE ID.")
            else:
                print("Please enter a valid CVE ID (e.g., CVE-2017-0144)")
        elif choice == "2":
            break
        else:
            print("Invalid input. Please enter a valid choice (1 or 2).")

    driver.quit()


if __name__ == "__main__":
    main()
